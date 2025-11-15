import os
import re
import shutil
from typing import List, Dict, Optional

from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import AzureOpenAIEmbeddings

# ==============================
#       LOAD ENV + CONFIG
# ==============================

load_dotenv()

# --- Azure OpenAI Config ---
AZURE_EMBEDDING_KEY = os.getenv("AZURE_OPENAI_EMBEDDING_API_KEY")
AZURE_EMBEDDING_ENDPOINT = os.getenv("AZURE_OPENAI_EMBEDDING_ENDPOINT")
AZURE_EMBEDDING_MODEL_NAME = os.getenv("AZURE_OPENAI_EMBED_MODEL")
AZURE_API_VERSION = "2023-05-15"

if not all([AZURE_EMBEDDING_KEY, AZURE_EMBEDDING_ENDPOINT, AZURE_EMBEDDING_MODEL_NAME]):
    raise ValueError("Azure OpenAI embedding environment variables not found.")

# --- File and Chroma Config ---
DOCUMENTS_DIR = "documents"
LAW_FILE = os.path.join(DOCUMENTS_DIR, "Law-36-2024-QH15.docx")
DECREE_FILE = os.path.join(DOCUMENTS_DIR, "1682024NĐ-CP.docx")

CHROMA_DB_DIR = "chroma_db"
COLLECTION_NAME = "traffic_law_2024"

# --- Regex for parsing ---
ARTICLE_RE = re.compile(r"^Điều\s+(\d+)\.?\s*(.*)", re.IGNORECASE)
CLAUSE_RE = re.compile(r"^(\d+)\.\s*(.*)")


# ==============================
#   CUSTOM CHUNKING LOGIC
# ==============================

def chunk_docx_by_article_clause(
    file_path: str,
    source_tag: str,
) -> List[Document]:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
    except Exception as e:
        print(f"Error loading DOCX file {file_path}: {e}")
        return []

    chunks: List[Document] = []
    current_article_number: Optional[int] = None
    current_article_title: str = ""
    current_clause_number: Optional[int] = None
    current_clause_lines: List[str] = []

    def flush_clause():
        if (
            current_article_number is not None
            and current_clause_number is not None
            and current_clause_lines
        ):
            content = "\n".join(current_clause_lines).strip()
            if not content:
                return

            metadata = {
                "source": source_tag,
                "source_file": os.path.basename(file_path),
                "article_number": current_article_number,
                "article_title": current_article_title.strip() or f"Điều {current_article_number}",
                "clause_number": current_clause_number,
            }
            chunk_doc = Document(page_content=content, metadata=metadata)
            chunks.append(chunk_doc)

    for text in paragraphs:
        text = text.strip()
        if not text:
            continue

        art_match = ARTICLE_RE.match(text)
        if art_match:
            flush_clause()
            current_clause_lines.clear()
            current_article_number = int(art_match.group(1))
            current_article_title = art_match.group(2) or ""
            current_clause_number = None
            continue

        clause_match = CLAUSE_RE.match(text)
        if clause_match and current_article_number is not None:
            flush_clause()
            current_clause_lines.clear()
            current_clause_number = int(clause_match.group(1))
            remainder = clause_match.group(2).strip()
            if remainder:
                current_clause_lines.append(remainder)
            continue

        if current_article_number is not None:
            if current_clause_number is None:
                current_clause_number = 0
                current_clause_lines.clear()
            current_clause_lines.append(text)

    flush_clause()
    return chunks


# ==============================
#   MAIN INGEST PIPELINE
# ==============================

def ingest_data():
    if os.path.exists(CHROMA_DB_DIR):
        print(f"Removing existing ChromaDB directory: {CHROMA_DB_DIR}")
        shutil.rmtree(CHROMA_DB_DIR)

    print("Starting document chunking...")
    law_chunks = chunk_docx_by_article_clause(LAW_FILE, "law_36_2024")
    print(f"  → Found {len(law_chunks)} chunks in {os.path.basename(LAW_FILE)}")
    decree_chunks = chunk_docx_by_article_clause(DECREE_FILE, "nd_168_2024")
    print(f"  → Found {len(decree_chunks)} chunks in {os.path.basename(DECREE_FILE)}")
    
    all_chunks = law_chunks + decree_chunks
    if not all_chunks:
        print("No chunks were created. Aborting ingestion.")
        return

    for i, chunk in enumerate(all_chunks):
        meta = chunk.metadata
        chunk.metadata["id"] = f"{meta['source']}-art{meta['article_number']}-clause{meta['clause_number']}-{i}"

    print("\nInitializing embeddings and ChromaDB...")
    embeddings = AzureOpenAIEmbeddings(
        azure_deployment=AZURE_EMBEDDING_MODEL_NAME,
        model=AZURE_EMBEDDING_MODEL_NAME,
        azure_endpoint=AZURE_EMBEDDING_ENDPOINT,
        api_key=AZURE_EMBEDDING_KEY,
        api_version=AZURE_API_VERSION,
        dimensions=512  #  OPTIMIZATION: Reduce embedding dimension
    )
    
    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=embeddings,
        persist_directory=CHROMA_DB_DIR,
    )
    
    print(f"Ingesting {len(all_chunks)} chunks into ChromaDB collection '{COLLECTION_NAME}'...")
    ids = [chunk.metadata["id"] for chunk in all_chunks]
    vectorstore.add_documents(documents=all_chunks, ids=ids)
    
    print("\n==============================")
    print("Ingestion completed successfully!")
    print(f"ChromaDB is persisted in: {os.path.abspath(CHROMA_DB_DIR)}")
    print(f"Total documents in collection: {vectorstore._collection.count()}")
    print("==============================\n")


if __name__ == "__main__":
    # This script requires 'python-docx' which is not in the main requirements.txt
    # You should have it installed locally to run this.
    try:
        import docx
    except ImportError:
        print("Error: 'python-docx' is not installed.")
        print("Please run 'pip install python-docx' to install it locally before running ingestion.")
        exit(1)
        
    ingest_data()
